import string
from bs4 import BeautifulSoup
import urllib.request
from docx import Document
from docx.shared import Inches

document = Document()


for i in range(1, 10) :

    uri = "https://www.jamaran.news/rbooks/read/%D9%86%D9%87%D8%AC%20%D8%A7%D9%84%D8%A8%D9%84%D8%A7%D8%BA%D9%87?type=ghesar&id=" + str(i)

    with urllib.request.urlopen(uri) as url:

        content = url.read()
        soup = BeautifulSoup(content, 'html.parser')

        wrapper = soup.find(
            "div", attrs={"class": "nahjol-content"}
        ).find('div', attrs={'class': 'nahjol-item'})


        title = wrapper.find(
            'h1', attrs={'class': 'center fb'}).contents[0].strip()

        document.add_heading(title, 0)

        divs = wrapper.findAll("div")
        for div in divs :
            paragraphs = div.findAll('p')
            for p in paragraphs :
                if (len(p) != 0) :

                    body = p.contents[0].strip()
                    p = document.add_paragraph(body)

    document.add_page_break()
    print(str(i) + " form 100")

document.save('demo.docx')